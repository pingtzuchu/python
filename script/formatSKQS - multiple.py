import os
import re
import docx
source=r"C:\Users\Ping-tzu 2 Chu A1\Google 雲端硬碟\Cloud Drive\\中國"
for root, dirs, filenames in os.walk(source):
    for f in filenames:
        if (f[-4:]=="docx") and (not f[0]=="~"):
            fullname=os.path.join(root,f)
            print (fullname)
            document=docx.Document(fullname)
            if document.paragraphs[0].text[0:2]=="##":
                newDoc="""<?xml version="1.0" encoding="UTF-8"?>
<?xml-model href="http://www.tei-c.org/release/xml/tei/custom/schema/relaxng/tei_all.rng" type="application/xml" schematypens="http://relaxng.org/ns/structure/1.0"?>
<?xml-model href="http://www.tei-c.org/release/xml/tei/custom/schema/relaxng/tei_all.rng" type="application/xml"
	schematypens="http://purl.oclc.org/dsdl/schematron"?>
<TEI xmlns="http://www.tei-c.org/ns/1.0">
  <teiHeader>
      <fileDesc>
         <titleStmt>
            <title>Title</title>
         </titleStmt>
         <publicationStmt>
            <p>Publication Information</p>
         </publicationStmt>
         <sourceDesc>
            <p>Information about the source</p>
         </sourceDesc>
      </fileDesc>
  </teiHeader>
  <text>
      <body>
                """ #newXMLfile with several quan <div>s
                quanTitle=[] #Title for Quan first level of div
                xmlLevel=0 #div level
                pageNumber=1 #page number
                lineNumber=1 #line number line8 add PageAB line16 add PageNumber
                lineBreak=0 #whether the line is broken
                hi=re.compile(r"(　[^\s　<>]+)\s([^\s　<>]*?)　", re.A)
                hi2=re.compile(r"　([^\s　<>]+)\s([^\s　<>]*?)$", re.A)
                for para in document.paragraphs: #deal with each paragraph
                    if para.text[0:2]=="##": # test whether begins with "##"
                        pageTitle=para.text.split("/")
                        if len(pageTitle)==1: #if only "##" => end of file
                            break
                        quan=pageTitle[len(pageTitle)-1]
                        lineNumber=0 #set lineNumber to 0
                        if not (quan in quanTitle): #test whether it is a new quan title
                            quanTitle.append(quan) #add the title into the quan title list
                            if not xmlLevel==0:
                                newDoc += "<pb n='"+str(pageNumber)+"b'/></p>\n" + "</div>" * xmlLevel
                            newDoc += "\n<div><head>" + quan + "</head>\n" #add in new quan title
                            pageNumber = 1
                            xmlLevel = 1
                        else:
                            newDoc += "<pb n='"+str(pageNumber)+"b'/></p>\n"
                            pageNumber += 1
                        newDoc += "<p>"
                    else:
                        if len(para.text)==0: # empty line
                            if lineBreak==0: # in not linebreak
                                lineNumber += 1
                        else: # not empty line
                            if lineBreak == 0:
                                text=para.text
                                while re.search(hi, text):
                                    text = re.sub(hi, r"<hi>\1\2</hi>　", text) # hi
                                if re.search(hi2, text): #hi2 first line of broken line
                                    text=re.sub(hi2, r"<hi>\1\2", text)
                                    lineBreak = 1
                                    lineNumber =  lineNumber-1
                                if re.match(r"\A　　　([^　])", text): #heading level 3
                                    if xmlLevel==3:
                                        text = re.sub(r"\A　　　([^　])", r"</p></div>\n<div><head>\1", text)
                                    else:
                                        text = re.sub(r"\A　　　([^　])", r"</p>\n<div><head>\1", text)
                                        xmlLevel=3
                                    text += "</head>\n<p>"
                                if re.match(r"\A　　([^　])", text): #heading level 2
                                    if xmlLevel==2:
                                        text = re.sub(r"\A　　([^　])", r"</p></div>\n<div><head>\1", text)
                                    elif xmlLevel==1:
                                        xmlLevel=2
                                        text = re.sub(r"\A　　([^　])", r"</p><div><head>\1", text)
                                    elif xmlLevel==3:
                                        xmlLevel=2
                                        text = re.sub(r"\A　　([^　])", r"</p></div></div>\n<div><head>\1", text)
                                    text += "</head>\n<p>"
                                newDoc += text+"\n"
                                lineNumber  += 1
                            else: #if there is a lineBreak
                                lineBreak = 0
                                lineNumber += 1
                                newDoc += para.text + "</hi>\n"
                        if lineNumber == 9:
                            newDoc += "<pb n='"+str(pageNumber)+"a'/>"
                newDoc += "<pb n='"+str(pageNumber)+"b'/></p>\n" + "</div>" * (xmlLevel+1) + "</body></text></TEI>"
                newName=r"C:\Users\Ping-tzu 2 Chu A1\Google 雲端硬碟\Cloud Drive\程式資料\Oxygen\ChinaXML" + fullname[58:-4]+"xml"
                fw = open(newName, "w", encoding="utf-8")
                fw.write(newDoc)
                fw.close
                os.remove(fullname)
